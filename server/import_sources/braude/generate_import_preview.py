#!/usr/bin/env python3
"""
DRY RUN preview generator for Braude import sources.
Does NOT write to Firebase.
"""
import json
import os
import re
import sys
import traceback
from datetime import date, datetime, time
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

SCRIPT_DIR = Path(__file__).resolve().parent
SERVER_DIR = SCRIPT_DIR.parent.parent
OUTPUT_PATH = SCRIPT_DIR / "import-preview.json"

# ---------------------------------------------------------------------------
# Shared helpers (mirrors yearbook_parser.py / labs_parser.py)
# ---------------------------------------------------------------------------

def normalize(s):
    if not s:
        return ""
    return re.sub(r"\s+", " ", str(s).replace("\u00A0", " ")).strip()


def is_course_code(text):
    return re.fullmatch(r"\d{5,6}", text or "") is not None


def safe_cell(row, idx, dash_as_zero=False):
    if idx is None or idx >= len(row.cells):
        return None
    v = normalize(row.cells[idx].text)
    if v == "":
        return None
    if v == "-":
        return 0 if dash_as_zero else None
    if re.fullmatch(r"\d+(\.\d+)?", v):
        return float(v) if "." in v else int(v)
    return None


def safe_hours(value):
    if isinstance(value, (int, float)) and 0 <= value <= 10:
        return value
    return None


def find_col(headers, *needles):
    for i, h in enumerate(headers):
        for n in needles:
            if n in h:
                return i
    return None


def find_relation_col(headers):
    for i, h in enumerate(headers):
        if "קדם" in h or "צמוד" in h:
            return i
    return None


def extract_relations_from_docx_cell(cell, course_name_map):
    relations = []
    for p in cell.paragraphs:
        line_text = normalize(p.text)
        if not line_text:
            continue
        codes = re.findall(r"\b\d{5,6}\b", line_text)
        if not codes:
            continue
        is_underlined = any(run.font.underline for run in p.runs)
        rel_type = "COREQUISITE" if is_underlined else "PREREQUISITE"
        for code in codes:
            relations.append({
                "courseCode": code,
                "courseName": course_name_map.get(code),
                "type": rel_type,
            })
    uniq = {}
    for r in relations:
        uniq[r["courseCode"]] = r
    return list(uniq.values())


REQUIRED_LAB_HEADERS = {
    "staff": ["שם המרצה", "מרצה"],
    "group": ["קבוצת מעבדה", "קבוצה"],
    "time": ["שעה"],
    "day": ["יום"],
    "date": ["תאריך"],
    "sessionNo": ["מס' מע'", "מספר מעבדה", "מס׳ מע", "מס' מעבדה"],
    "sessionName": ["שם המקצוע", "שם הקורס"],
}


def build_lab_header(ws, row, max_col):
    cells = [normalize(ws.cell(row=row, column=c).value) for c in range(1, max_col + 1)]
    mapping = {}
    hits = 0
    for key, variants in REQUIRED_LAB_HEADERS.items():
        for i, txt in enumerate(cells, start=1):
            if txt and any(v in txt for v in variants):
                mapping[key] = i
                hits += 1
                break
    return mapping, hits


def extract_course_from_line(text):
    t = normalize(text)
    if not t:
        return None, None
    m = re.search(r"(.+?)\s*[-–]\s*(\d{4,6})$", t)
    if m:
        return m.group(2), normalize(m.group(1))
    m = re.search(r"^(\d{4,6})\s*[-–]\s*(.+)$", t)
    if m:
        return m.group(1), normalize(m.group(2))
    return None, None


def find_course_title_near(ws, header_row, max_col):
    for r in range(header_row - 1, max(1, header_row - 8), -1):
        line = " ".join(
            normalize(ws.cell(row=r, column=c).value)
            for c in range(1, max_col + 1)
            if ws.cell(row=r, column=c).value
        )
        code, name = extract_course_from_line(line)
        if code and name:
            return code, name
    return None, None


def json_serializable(obj):
    if isinstance(obj, (datetime, date, time)):
        return obj.isoformat()
    raise TypeError(f"Not serializable: {type(obj)}")


# ---------------------------------------------------------------------------
# File classification
# ---------------------------------------------------------------------------

YEARBOOK_KEYWORDS = {
    "required_courses_plan": ["קורסי חובה", "חובה"],
    "specialization_plan": ["התמחות", "מסלול"],
    "general_guidelines": ["הנחיות כלליות", "כללי"],
}

SEMESTER_RE = re.compile(r"סמסטר\s*['\u05f3]?\s*(\d+)", re.IGNORECASE)
SEMESTER_SHORT_RE = re.compile(r"(?:^|\s)([1-8])(?:\s|$)")


def classify_file(rel_path: str) -> dict:
    name = Path(rel_path).name
    lower = name.lower()
    ext = Path(name).suffix.lower()
    parent = Path(rel_path).parent.as_posix()

    info = {
        "fileName": name,
        "relativePath": rel_path.replace("\\", "/"),
        "extension": ext,
        "type": "unknown",
        "semesters": [],
        "yearbookHint": None,
        "compatibleParser": None,
        "parserNotes": [],
    }

    sem_match = SEMESTER_RE.search(name)
    if sem_match:
        info["semesters"] = [int(sem_match.group(1))]

    if parent == "yearbooks" or "yearbooks" in parent:
        info["type"] = "yearbook"
        info["compatibleParser"] = "yearbook_parser.py"
        if ext != ".docx":
            info["parserNotes"].append(f"Parser expects .docx, got {ext}")
        for subtype, keywords in YEARBOOK_KEYWORDS.items():
            if any(kw in name for kw in keywords):
                info["type"] = subtype
                break
        info["yearbookHint"] = suggest_yearbook_id(name, info["type"])
        return info

    if parent == "labs" or "labs" in parent:
        info["type"] = "labs"
        info["compatibleParser"] = "labs_parser.py"
        if ext != ".xlsx":
            info["parserNotes"].append(f"Parser expects .xlsx, got {ext}")
        if not info["semesters"]:
            m = re.search(r"סמסטר\s*(\d+)", name)
            if m:
                info["semesters"] = [int(m.group(1))]
        return info

    if parent == "registration_guidelines" or "registration" in parent:
        info["type"] = "registration_guidelines"
        info["compatibleParser"] = None
        info["parserNotes"].append("No automated parser — heuristic extraction only")
        if ext == ".doc":
            info["parserNotes"].append(".doc format not supported by python-docx; manual conversion needed")
        if not info["semesters"]:
            m = re.search(r"(?:סמסטר|סמ)\s*['\u05f3]?\s*(\d+)", name, re.I)
            if m:
                info["semesters"] = [int(m.group(1))]
            else:
                m2 = re.search(r"\b([1-8])\b", name)
                if m2:
                    info["semesters"] = [int(m2.group(1))]
        return info

    if any(kw in name for kw in ["הנחיות כלליות", "כללי"]):
        info["type"] = "general_guidelines"
        info["parserNotes"].append("No parser for general guidelines")
        return info

    if ext == ".xlsx" and any(kw in name for kw in ["פריסה", "קורסים"]):
        info["type"] = "required_courses_plan"
        info["parserNotes"].append("XLSX course plan — no dedicated parser")
        return info

    return info


def suggest_yearbook_id(filename: str, file_type: str) -> dict:
    """Suggest yearbookId and label from filename — not invented, derived from file name."""
    name = Path(filename).stem
    slug = re.sub(r"[^\w\u0590-\u05FF]+", "_", name).strip("_").lower()
    slug = re.sub(r"_+", "_", slug)

    label = name
    yearbook_id = slug[:80] if slug else "unknown"

    track_hint = None
    if "מולקולרית" in name or "תרופתית" in name:
        track_hint = "מולקולרית-תרופתית"
    elif "מזון" in name or "הסביבה" in name:
        track_hint = "מזון והסביבה"
    elif "ביוטכנולוגיה" in name:
        track_hint = "ביוטכנולוגיה"

    return {
        "suggestedYearbookId": yearbook_id,
        "suggestedLabel": label,
        "trackHint": track_hint,
        "source": "filename",
    }


# ---------------------------------------------------------------------------
# Yearbook preview (dry-run of yearbook_parser.process_docx)
# ---------------------------------------------------------------------------

def preview_yearbook(file_path: Path, meta: dict) -> dict:
    hint = meta.get("yearbookHint") or {}
    result = {
        "sourceFile": meta["relativePath"],
        "fileType": meta["type"],
        "yearbookId": hint.get("suggestedYearbookId"),
        "label": hint.get("suggestedLabel"),
        "trackHint": hint.get("trackHint"),
        "semesters": [],
        "courses": [],
        "courseCount": 0,
        "relationCount": 0,
        "parseStatus": "ok",
        "parseErrors": [],
        "compatibleParser": "yearbook_parser.py",
    }

    if meta["extension"] != ".docx":
        result["parseStatus"] = "skipped"
        result["parseErrors"].append("Only .docx supported")
        return result

    try:
        doc = Document(str(file_path))
    except Exception as e:
        result["parseStatus"] = "error"
        result["parseErrors"].append(str(e))
        return result

    table_map = {t._element: t for t in doc.tables}
    current_sem = None
    created_semesters = set()
    course_name_map = {}

    for block in doc.element.body:
        if block.tag.endswith("p"):
            text = normalize("".join(t.text for t in block.xpath(".//w:t")))
            m = re.search(r"סמסטר\s*([1-8])", text)
            if m:
                current_sem = int(m.group(1))
                if current_sem not in created_semesters:
                    result["semesters"].append({"semesterNumber": current_sem})
                    created_semesters.add(current_sem)

        elif block.tag.endswith("tbl") and current_sem:
            table = table_map.get(block)
            if not table or not table.rows:
                continue
            headers = [normalize(c.text) for c in table.rows[0].cells]
            if "שם הקורס" not in " ".join(headers):
                continue

            code_i, name_i = 0, 1
            lec_i = find_col(headers, "הרצאה", "ה")
            prac_i = find_col(headers, "תרגול", "ת")
            lab_i = find_col(headers, "מעבדה", "מ")
            cred_i = find_col(headers, 'נ"ז', "נקודות", "נ")
            rel_i = find_relation_col(headers)

            for row in table.rows[1:]:
                code = normalize(row.cells[code_i].text)
                name = normalize(row.cells[name_i].text)
                if not is_course_code(code):
                    continue
                if name:
                    course_name_map[code] = name

                relations = []
                if rel_i is not None:
                    relations = extract_relations_from_docx_cell(row.cells[rel_i], course_name_map)
                    for r in relations:
                        if r.get("courseName") is None and r["courseCode"] in course_name_map:
                            r["courseName"] = course_name_map[r["courseCode"]]

                course = {
                    "semesterNumber": current_sem,
                    "courseCode": code,
                    "courseName": name,
                    "lectureHours": safe_hours(safe_cell(row, lec_i, True)),
                    "practiceHours": safe_hours(safe_cell(row, prac_i, True)),
                    "labHours": safe_hours(safe_cell(row, lab_i, True)),
                    "credits": safe_cell(row, cred_i),
                    "relations": relations,
                }
                result["courses"].append(course)
                result["relationCount"] += len(relations)

    result["courseCount"] = len(result["courses"])
    if result["courseCount"] == 0:
        result["parseStatus"] = "warning"
        result["parseErrors"].append("No courses extracted — check document structure")
    return result


# ---------------------------------------------------------------------------
# Labs preview (dry-run of labs_parser.parse_workbook)
# ---------------------------------------------------------------------------

def preview_labs(file_path: Path, meta: dict) -> dict:
    semester = meta["semesters"][0] if meta["semesters"] else None
    year_hint = re.search(r"20\d{2}", meta["fileName"])
    year_id = f"labs_{year_hint.group()}" if year_hint else "labs_unknown_year"

    result = {
        "sourceFile": meta["relativePath"],
        "fileType": "labs",
        "semester": semester,
        "suggestedYearId": year_id,
        "suggestedYearLabel": year_hint.group() if year_hint else None,
        "courses": [],
        "labRecords": [],
        "courseCount": 0,
        "labRecordCount": 0,
        "parseStatus": "ok",
        "parseErrors": [],
        "compatibleParser": "labs_parser.py",
    }

    if meta["extension"] != ".xlsx":
        result["parseStatus"] = "skipped"
        result["parseErrors"].append("Only .xlsx supported")
        return result

    try:
        wb = load_workbook(file_path, data_only=True)
    except Exception as e:
        result["parseStatus"] = "error"
        result["parseErrors"].append(str(e))
        return result

    courses_map = {}

    for ws in wb.worksheets:
        r = 1
        while r <= ws.max_row:
            header, hits = build_lab_header(ws, r, ws.max_column)
            if hits < 5:
                r += 1
                continue

            course_code, course_name = find_course_title_near(ws, r, ws.max_column)
            if not course_code:
                r += 1
                continue

            courses_map.setdefault(
                course_code,
                {"courseCode": str(course_code), "courseName": course_name, "labs": []},
            )

            rr = r + 1
            while rr <= ws.max_row:
                date_val = normalize(ws.cell(rr, header.get("date", 0)).value)
                day_val = normalize(ws.cell(rr, header.get("day", 0)).value)
                session_val = ""
                if header.get("sessionNo"):
                    session_val = normalize(ws.cell(rr, header["sessionNo"]).value)
                if not session_val and header.get("sessionName"):
                    session_val = normalize(ws.cell(rr, header["sessionName"]).value)

                if not session_val and not date_val and not day_val:
                    break

                lab = {
                    "courseCode": str(course_code),
                    "courseName": course_name,
                    "semester": semester,
                    "yearId": year_id,
                    "session": session_val,
                    "date": date_val,
                    "day": day_val,
                    "group": normalize(ws.cell(rr, header.get("group", 0)).value),
                    "time": normalize(ws.cell(rr, header.get("time", 0)).value),
                    "staff": [],
                }
                if header.get("staff"):
                    staff = normalize(ws.cell(rr, header["staff"]).value)
                    if staff:
                        lab["staff"] = [staff]

                courses_map[course_code]["labs"].append(lab)
                result["labRecords"].append(lab)
                rr += 1
            r = rr

    result["courses"] = list(courses_map.values())
    result["courseCount"] = len(result["courses"])
    result["labRecordCount"] = len(result["labRecords"])

    if result["labRecordCount"] == 0:
        result["parseStatus"] = "warning"
        result["parseErrors"].append("No lab records extracted")
    if semester is None:
        result["parseStatus"] = "warning"
        result["parseErrors"].append("Semester not detected from filename — required for labs_parser.py")
    return result


# ---------------------------------------------------------------------------
# Registration guidelines heuristic preview
# ---------------------------------------------------------------------------

EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(r"0\d{1,2}[-\s]?\d{3,4}[-\s]?\d{3,4}")
URL_RE = re.compile(r"https?://\S+|www\.\S+")
DATE_RE = re.compile(r"\b(\d{1,2})[./](\d{1,2})[./](\d{2,4})\b")
TIME_RE = re.compile(r"\b(\d{1,2}):(\d{2})\b")


def docx_paragraphs_text(file_path: Path) -> list[str]:
    doc = Document(str(file_path))
    return [normalize(p.text) for p in doc.paragraphs if normalize(p.text)]


def preview_registration_guidelines(file_path: Path, meta: dict) -> dict:
    semester = meta["semesters"][0] if meta["semesters"] else None
    result = {
        "sourceFile": meta["relativePath"],
        "fileType": "registration_guidelines",
        "semesterNumber": semester,
        "docId": f"semester_{semester}" if semester else None,
        "title": None,
        "term": None,
        "audience": {"cohortText": None, "creditsRuleText": None},
        "registrationWindow": {"date": None, "from": None, "to": None, "rawText": []},
        "keyRules": [],
        "links": [],
        "contacts": {
            "registrationSupport": [],
            "mentors": [],
            "academicAdvisors": [],
            "exemptions": [],
            "labs": [],
        },
        "extractedAdvisors": [],
        "parseStatus": "ok",
        "parseErrors": [],
        "extractionWarnings": [],
        "compatibleParser": None,
    }

    if meta["extension"] == ".doc":
        result["parseStatus"] = "skipped"
        result["parseErrors"].append("Legacy .doc format — convert to .docx for automated extraction")
        return result

    if meta["extension"] != ".docx":
        result["parseStatus"] = "skipped"
        result["parseErrors"].append(f"Unsupported extension: {meta['extension']}")
        return result

    try:
        paragraphs = docx_paragraphs_text(file_path)
    except Exception as e:
        result["parseStatus"] = "error"
        result["parseErrors"].append(str(e))
        return result

    full_text = "\n".join(paragraphs)

    # Title: first substantial line or line with "הנחיות רישום"
    for p in paragraphs[:5]:
        if "הנחיות" in p and len(p) > 10:
            result["title"] = p
            break
    if not result["title"] and paragraphs:
        result["title"] = paragraphs[0]

    # Term A/B
    term_m = re.search(r"סמסטר\s*[ABאב]|מחצית\s*[ABאב]|[\s(]([ABאב])[\s)]", full_text)
    if term_m:
        t = term_m.group(1) if term_m.lastindex else term_m.group(0)
        result["term"] = t[-1] if t else None

    # Audience / cohort
    for p in paragraphs:
        if any(kw in p for kw in ["שנתון", "קהל יעד", "לסטודנטים", "ביוטכנולוגיה"]):
            if len(p) > 15 and not result["audience"]["cohortText"]:
                result["audience"]["cohortText"] = p

    # Credits rules
    for p in paragraphs:
        if 'נ"ז' in p or "נקודות זכות" in p or "נק״ז" in p:
            result["audience"]["creditsRuleText"] = p
            break

    # Registration window
    for p in paragraphs:
        if any(kw in p for kw in ["רישום", "פתיחת הרישום", "חלון רישום", "שעות הרישום"]):
            result["registrationWindow"]["rawText"].append(p)
            dates = DATE_RE.findall(p)
            times = TIME_RE.findall(p)
            if dates and not result["registrationWindow"]["date"]:
                d, m, y = dates[0]
                y = y if len(y) == 4 else f"20{y}"
                result["registrationWindow"]["date"] = f"{y}-{m.zfill(2)}-{d.zfill(2)}"
            if len(times) >= 1 and not result["registrationWindow"]["from"]:
                result["registrationWindow"]["from"] = f"{times[0][0].zfill(2)}:{times[0][1]}"
            if len(times) >= 2 and not result["registrationWindow"]["to"]:
                result["registrationWindow"]["to"] = f"{times[1][0].zfill(2)}:{times[1][1]}"

    # Key rules — bullet-like or numbered lines with imperative language
    rule_keywords = ["חובה", "אסור", "יש ל", "נא ל", "שימו לב", "חשוב"]
    for p in paragraphs:
        if len(p) < 20:
            continue
        if any(kw in p for kw in rule_keywords):
            if p not in result["registrationWindow"]["rawText"]:
                result["keyRules"].append({"code": f"RULE_{len(result['keyRules']) + 1}", "text": p})

    # Links
    for url in URL_RE.findall(full_text):
        result["links"].append({"label": url, "url": url})

    # Contacts — emails with surrounding context
    for i, p in enumerate(paragraphs):
        emails = EMAIL_RE.findall(p)
        phones = PHONE_RE.findall(p)
        if not emails and not phones:
            continue

        person = {"name": None, "role": None, "email": emails[0] if emails else None, "phone": phones[0] if phones else None, "sourceLine": p}

        # Try to extract name from same line (before email)
        if emails:
            before = p.split(emails[0])[0].strip(" :-–—")
            if 3 < len(before) < 60:
                person["name"] = before

        lower = p.lower()
        if "יועץ" in p or "יועצת" in p:
            advisor = {
                "name": person["name"],
                "email": person["email"],
                "assignment": {"lastNameFrom": None, "lastNameTo": None, "track": None},
                "sourceLine": p,
            }
            alpha = re.search(r"([א-ת])\s*[-–]\s*([א-ת])", p)
            if alpha:
                advisor["assignment"]["lastNameFrom"] = alpha.group(1)
                advisor["assignment"]["lastNameTo"] = alpha.group(2)
            result["contacts"]["academicAdvisors"].append(advisor)
            result["extractedAdvisors"].append(advisor)
        elif "מנטור" in p or "מלווה" in p:
            result["contacts"]["mentors"].append(person)
        elif "מעבדה" in p or "מעבדות" in p:
            person["howToContact"] = p
            result["contacts"]["labs"].append(person)
        elif "פטור" in p or "חריג" in p:
            result["contacts"]["exemptions"].append(person)
        else:
            result["contacts"]["registrationSupport"].append(person)

    if semester is None:
        result["extractionWarnings"].append("Semester not detected from filename")
        result["parseStatus"] = "warning"

    if not result["keyRules"] and not result["registrationWindow"]["rawText"]:
        result["extractionWarnings"].append("Limited structured content extracted — manual review recommended")

    return result


# ---------------------------------------------------------------------------
# Main scan
# ---------------------------------------------------------------------------

def scan_all():
    preview = {
        "generatedAt": datetime.now().isoformat(),
        "mode": "DRY_RUN",
        "sourceRoot": str(SCRIPT_DIR.relative_to(SERVER_DIR)).replace("\\", "/"),
        "yearbooks": [],
        "labs": [],
        "registrationGuidelines": [],
        "advisors": [],
        "warnings": [],
        "unknownFiles": [],
        "fileInventory": [],
        "summary": {},
    }

    all_files = []
    for dp, _, fns in os.walk(SCRIPT_DIR):
        for fn in fns:
            if fn in ("import-preview.json", "generate_import_preview.py"):
                continue
            full = Path(dp) / fn
            rel = full.relative_to(SCRIPT_DIR).as_posix()
            all_files.append(rel)

    all_files.sort()

    for rel in all_files:
        full_path = SCRIPT_DIR / rel
        meta = classify_file(rel)
        meta["fileSizeBytes"] = full_path.stat().st_size
        preview["fileInventory"].append(meta)

        try:
            if meta["type"] in ("yearbook", "required_courses_plan", "specialization_plan", "general_guidelines"):
                if meta["type"] == "yearbook":
                    entry = preview_yearbook(full_path, meta)
                else:
                    # Same parser may work for course plans in docx
                    if meta["extension"] == ".docx":
                        entry = preview_yearbook(full_path, meta)
                        entry["fileType"] = meta["type"]
                    else:
                        entry = {
                            "sourceFile": rel,
                            "fileType": meta["type"],
                            "parseStatus": "skipped",
                            "parseErrors": meta["parserNotes"] or ["No parser for this format"],
                        }
                preview["yearbooks"].append(entry)
                if entry.get("parseStatus") in ("warning", "error", "skipped"):
                    preview["warnings"].append({
                        "file": rel,
                        "category": "yearbooks",
                        "status": entry.get("parseStatus"),
                        "messages": entry.get("parseErrors", []) + meta.get("parserNotes", []),
                    })

            elif meta["type"] == "labs":
                entry = preview_labs(full_path, meta)
                preview["labs"].append(entry)
                if entry.get("parseStatus") in ("warning", "error", "skipped"):
                    preview["warnings"].append({
                        "file": rel,
                        "category": "labs",
                        "status": entry.get("parseStatus"),
                        "messages": entry.get("parseErrors", []) + meta.get("parserNotes", []),
                    })

            elif meta["type"] == "registration_guidelines":
                entry = preview_registration_guidelines(full_path, meta)
                preview["registrationGuidelines"].append(entry)
                for adv in entry.get("extractedAdvisors", []):
                    preview["advisors"].append({
                        "sourceFile": rel,
                        "semesterNumber": entry.get("semesterNumber"),
                        **adv,
                    })
                if entry.get("parseStatus") in ("warning", "error", "skipped"):
                    preview["warnings"].append({
                        "file": rel,
                        "category": "registrationGuidelines",
                        "status": entry.get("parseStatus"),
                        "messages": entry.get("parseErrors", []) + entry.get("extractionWarnings", []) + meta.get("parserNotes", []),
                    })

            else:
                preview["unknownFiles"].append({
                    "fileName": meta["fileName"],
                    "relativePath": rel,
                    "extension": meta["extension"],
                    "reason": "Could not classify file type",
                })

        except Exception as e:
            preview["warnings"].append({
                "file": rel,
                "category": "exception",
                "status": "error",
                "messages": [str(e), traceback.format_exc()],
            })

    # Summary
    type_counts = {}
    for m in preview["fileInventory"]:
        type_counts[m["type"]] = type_counts.get(m["type"], 0) + 1

    courses_by_yearbook = {}
    for yb in preview["yearbooks"]:
        yid = yb.get("yearbookId") or yb.get("sourceFile")
        courses_by_yearbook[yid] = courses_by_yearbook.get(yid, 0) + yb.get("courseCount", 0)

    preview["summary"] = {
        "totalFiles": len(all_files),
        "filesByType": type_counts,
        "yearbookFiles": len(preview["yearbooks"]),
        "labsFiles": len(preview["labs"]),
        "registrationGuidelineFiles": len(preview["registrationGuidelines"]),
        "unknownFileCount": len(preview["unknownFiles"]),
        "warningCount": len(preview["warnings"]),
        "totalCoursesExtracted": sum(yb.get("courseCount", 0) for yb in preview["yearbooks"]),
        "coursesByYearbook": courses_by_yearbook,
        "totalLabRecords": sum(lb.get("labRecordCount", 0) for lb in preview["labs"]),
        "totalLabCourses": sum(lb.get("courseCount", 0) for lb in preview["labs"]),
        "registrationGuidelinesWithContent": sum(
            1 for rg in preview["registrationGuidelines"] if rg.get("parseStatus") == "ok"
        ),
        "advisorsExtracted": len(preview["advisors"]),
        "readyForFirebaseWrite": False,
        "readinessNotes": [],
    }

    # Readiness assessment
    notes = preview["summary"]["readinessNotes"]
    blocking = False

    if preview["unknownFiles"]:
        blocking = True
        notes.append(f"{len(preview['unknownFiles'])} unclassified file(s)")

    skipped_regs = [rg for rg in preview["registrationGuidelines"] if rg.get("parseStatus") == "skipped"]
    if skipped_regs:
        blocking = True
        notes.append(f"{len(skipped_regs)} registration guideline file(s) skipped (.doc or unsupported)")

    error_entries = [w for w in preview["warnings"] if w.get("status") == "error"]
    if error_entries:
        blocking = True
        notes.append(f"{len(error_entries)} file(s) with parse errors")

    zero_courses = [yb for yb in preview["yearbooks"] if yb.get("courseCount", 0) == 0 and yb.get("parseStatus") != "skipped"]
    if zero_courses:
        notes.append(f"{len(zero_courses)} yearbook file(s) extracted 0 courses — verify before import")

    zero_labs = [lb for lb in preview["labs"] if lb.get("labRecordCount", 0) == 0]
    if zero_labs:
        notes.append(f"{len(zero_labs)} labs file(s) extracted 0 records")

    no_semester_labs = [lb for lb in preview["labs"] if lb.get("semester") is None]
    if no_semester_labs:
        blocking = True
        notes.append("Some labs files missing semester — labs_parser requires semester parameter")

    if not blocking and preview["summary"]["totalCoursesExtracted"] > 0:
        preview["summary"]["readyForFirebaseWrite"] = True
        notes.append("Yearbooks and labs look parseable; registration guidelines need manual review of heuristic extraction")
    elif not blocking:
        notes.append("No blocking issues but little data extracted — review preview before write phase")

    return preview


def main():
    import argparse

    parser = argparse.ArgumentParser(description="Generate Braude import previews (dry-run).")
    parser.add_argument(
        "--registration-guidelines-only",
        action="store_true",
        help="Generate registration-guidelines-preview.json only",
    )
    args = parser.parse_args()

    if args.registration_guidelines_only:
        from registration_guidelines_parser import scan_all_registration_guidelines

        reg_dir = SCRIPT_DIR / "registration_guidelines"
        out_path = SCRIPT_DIR / "registration-guidelines-preview.json"
        preview = scan_all_registration_guidelines(reg_dir)
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(preview, f, ensure_ascii=False, indent=2, default=json_serializable)
        print(f"Wrote preview to {out_path}")
        print(json.dumps(preview["summary"], ensure_ascii=False, indent=2))
        return

    preview = scan_all()
    with open(OUTPUT_PATH, "w", encoding="utf-8") as f:
        json.dump(preview, f, ensure_ascii=False, indent=2, default=json_serializable)
    print(f"Wrote preview to {OUTPUT_PATH}")
    print(json.dumps(preview["summary"], ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
