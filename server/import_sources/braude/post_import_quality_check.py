#!/usr/bin/env python3
"""
Post-import quality check — READ ONLY.
Compares Firebase data with local source files.
Does NOT write to Firebase.
"""
from __future__ import annotations

import json
import os
import re
import sys
from collections import defaultdict
from datetime import date, datetime, time
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
import firebase_admin
from firebase_admin import credentials, firestore

SCRIPT_DIR = Path(__file__).resolve().parent
SERVER_DIR = SCRIPT_DIR.parent.parent
ENV_FILE = SERVER_DIR / ".env"
OUTPUT_PATH = SCRIPT_DIR / "post-import-quality-report.json"

YEARBOOK_SOURCES = {
    "tashpag": SCRIPT_DIR / "yearbooks" / "שנתון  תשפג מעודכן.docx",
    "tashpad": SCRIPT_DIR / "yearbooks" / "שנתון תשפד- עדכון התמחויות.docx",
    "tashpah": SCRIPT_DIR / "yearbooks" / "שנתון תשפה- מעודכן.docx",
    "tashpav": SCRIPT_DIR / "yearbooks" / "שנתון תשפו.docx",
}

LAB_SEMESTER_RE = re.compile(r"סמסטר\s*(\d+)")


def load_dotenv_readonly(env_path: Path) -> None:
    if not env_path.is_file():
        return
    for raw in env_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, _, value = line.partition("=")
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


def init_firebase():
    if not firebase_admin._apps:
        load_dotenv_readonly(ENV_FILE)
        cred = credentials.Certificate({
            "type": "service_account",
            "project_id": os.environ["FIREBASE_PROJECT_ID"],
            "client_email": os.environ["FIREBASE_CLIENT_EMAIL"],
            "private_key": os.environ["FIREBASE_PRIVATE_KEY"].replace("\\n", "\n"),
            "token_uri": "https://oauth2.googleapis.com/token",
        })
        firebase_admin.initialize_app(cred)
    return firestore.client()


def normalize(s):
    if not s:
        return ""
    return re.sub(r"\s+", " ", str(s).replace("\u00A0", " ")).strip()


def is_course_code(text):
    return re.fullmatch(r"\d{5,6}", text or "") is not None


def find_col(headers, *needles):
    for i, h in enumerate(headers):
        for n in needles:
            if n in h:
                return i
    return None


def find_relation_col(headers):
    for i, h in enumerate(headers):
        if "קדם" in h or "צמוד" in h:
            return i
    return None


def parse_credits_raw(cell_text):
    v = normalize(cell_text)
    if not v or v == "-":
        return None, v
    if re.fullmatch(r"\d+(\.\d+)?", v):
        return (float(v) if "." in v else int(v)), v
    return None, v


def extract_docx_courses_detailed(docx_path: Path) -> list[dict]:
    doc = Document(str(docx_path))
    table_map = {t._element: t for t in doc.tables}
    current_sem = None
    courses = []

    for block in doc.element.body:
        if block.tag.endswith("p"):
            text = normalize("".join(t.text for t in block.xpath(".//w:t")))
            m = re.search(r"סמסטר\s*([1-8])", text)
            if m:
                current_sem = int(m.group(1))
        elif block.tag.endswith("tbl") and current_sem:
            table = table_map.get(block)
            if not table or not table.rows:
                continue
            headers = [normalize(c.text) for c in table.rows[0].cells]
            header_joined = " ".join(headers)
            if "שם הקורס" not in header_joined:
                continue

            cred_i = find_col(headers, 'נ"ז', "נקודות", "נ")
            lec_i = find_col(headers, "הרצאה", "ה")
            prac_i = find_col(headers, "תרגול", "ת")
            lab_i = find_col(headers, "מעבדה", "מ")

            for row in table.rows[1:]:
                code = normalize(row.cells[0].text)
                name = normalize(row.cells[1].text)
                if not is_course_code(code):
                    continue

                raw_cred = row.cells[cred_i].text if cred_i is not None and cred_i < len(row.cells) else ""
                parsed_cred, raw_cred_norm = parse_credits_raw(raw_cred)

                courses.append({
                    "semester": current_sem,
                    "courseCode": code,
                    "courseName": name,
                    "creditsParsed": parsed_cred,
                    "creditsRaw": raw_cred_norm,
                    "credColIndex": cred_i,
                    "headers": headers,
                    "allCells": [normalize(c.text) for c in row.cells],
                })
    return courses


def inspect_test_yearbook(db) -> dict:
    ref = db.collection("yearbooks").document("test")
    doc = ref.get()
    result = {
        "exists": doc.exists,
        "documentId": "test",
        "data": doc.to_dict() if doc.exists else None,
        "requiredCoursesCount": 0,
        "semesterIds": [],
        "coursesCount": 0,
        "relationsCount": 0,
        "looksLikeEmptyTest": False,
        "deleteRecommendation": None,
    }
    if not doc.exists:
        result["deleteRecommendation"] = "not_found"
        return result

    sem_snap = ref.collection("requiredCourses").stream()
    for sem in sem_snap:
        result["requiredCoursesCount"] += 1
        result["semesterIds"].append(sem.id)
        for course in sem.reference.collection("courses").stream():
            result["coursesCount"] += 1
            for _ in course.reference.collection("relations").stream():
                result["relationsCount"] += 1

    result["looksLikeEmptyTest"] = (
        result["requiredCoursesCount"] == 0
        and result["coursesCount"] == 0
        and result["relationsCount"] == 0
    )
    if result["looksLikeEmptyTest"]:
        result["deleteRecommendation"] = "safe_to_delete_after_approval"
    elif result["coursesCount"] == 0:
        result["deleteRecommendation"] = "likely_stale_shell_delete_after_approval"
    else:
        result["deleteRecommendation"] = "do_not_delete_has_data"

    return result


def fetch_firebase_courses(db) -> list[dict]:
    rows = []
    for yb in db.collection("yearbooks").stream():
        yb_id = yb.id
        for sem in yb.reference.collection("requiredCourses").stream():
            for course in sem.reference.collection("courses").stream():
                c = course.to_dict() or {}
                rows.append({
                    "yearbookId": yb_id,
                    "semester": sem.id,
                    "semesterNumber": (sem.to_dict() or {}).get("semesterNumber"),
                    "courseCode": c.get("courseCode") or course.id,
                    "courseName": c.get("courseName"),
                    "credits": c.get("credits"),
                })
    return rows


def analyze_missing_credits(db) -> dict:
    fb_courses = fetch_firebase_courses(db)
    missing_fb = [c for c in fb_courses if c["credits"] is None and c["yearbookId"] != "test"]

    # Source analysis per imported yearbook
    source_by_key: dict[tuple, dict] = {}
    for yb_id, path in YEARBOOK_SOURCES.items():
        if not path.is_file():
            continue
        for row in extract_docx_courses_detailed(path):
            key = (yb_id, row["semester"], row["courseCode"])
            source_by_key[key] = {**row, "yearbookId": yb_id}

    report_rows = []
    unique_codes: set[str] = set()

    for fb in missing_fb:
        sem_num = fb["semesterNumber"]
        if sem_num is None and fb["semester"].startswith("semester_"):
            try:
                sem_num = int(fb["semester"].split("_")[1])
            except (IndexError, ValueError):
                sem_num = None

        key = (fb["yearbookId"], sem_num, fb["courseCode"])
        src = source_by_key.get(key)
        in_all_four = True
        src_states = {}
        for yb_id in YEARBOOK_SOURCES:
            k = (yb_id, sem_num, fb["courseCode"])
            s = source_by_key.get(k)
            src_states[yb_id] = {
                "creditsRaw": s["creditsRaw"] if s else None,
                "creditsParsed": s["creditsParsed"] if s else None,
            }
            if not s or s["creditsParsed"] is None:
                in_all_four = False

        diagnosis = "unknown"
        if src:
            if src["creditsParsed"] is not None:
                diagnosis = "parser_or_import_gap_source_has_value"
            elif src["creditsRaw"] in ("", "-", None):
                diagnosis = "genuinely_missing_in_source_docx"
            else:
                diagnosis = "parser_failed_non_numeric_source"
        else:
            diagnosis = "course_not_found_in_source_file"

        unique_codes.add(fb["courseCode"])
        report_rows.append({
            "courseCode": fb["courseCode"],
            "courseName": fb["courseName"],
            "yearbookId": fb["yearbookId"],
            "semester": fb["semester"],
            "semesterNumber": sem_num,
            "firebaseCredits": fb["credits"],
            "sourceCreditsRaw": src["creditsRaw"] if src else None,
            "sourceCreditsParsed": src["creditsParsed"] if src else None,
            "missingInAllFourYearbooks": in_all_four,
            "perYearbookSource": src_states,
            "diagnosis": diagnosis,
            "sourceHeadersSample": src["headers"] if src else None,
            "credColIndex": src["credColIndex"] if src else None,
        })

    # Unique course codes missing credits (dedupe across yearbooks)
    by_code: dict[str, list] = defaultdict(list)
    for r in report_rows:
        by_code[r["courseCode"]].append(r)

    unique_missing = []
    for code, items in sorted(by_code.items()):
        diagnoses = {i["diagnosis"] for i in items}
        unique_missing.append({
            "courseCode": code,
            "courseName": items[0]["courseName"],
            "occurrences": len(items),
            "yearbooks": sorted({i["yearbookId"] for i in items}),
            "missingInAllFourYearbooks": all(i["missingInAllFourYearbooks"] for i in items),
            "diagnoses": sorted(diagnoses),
            "sourceCreditsRawSample": items[0].get("sourceCreditsRaw"),
        })

    return {
        "firebaseMissingCreditsCount": len(missing_fb),
        "uniqueCourseCodesMissingCredits": len(unique_codes),
        "rows": report_rows,
        "uniqueSummary": unique_missing,
    }


def merged_cell_value(ws, row, col):
    """Return value respecting merged cells (top-left anchor)."""
    cell = ws.cell(row=row, column=col)
    for rng in ws.merged_cells.ranges:
        if cell.coordinate in rng:
            return ws.cell(row=rng.min_row, column=rng.min_col).value
    return cell.value


def norm(x):
    if x is None:
        return ""
    if isinstance(x, (datetime, date, time)):
        return x.isoformat() if hasattr(x, "isoformat") else str(x)
    return re.sub(r"\s+", " ", str(x)).strip()


def build_lab_header(ws, row, max_col):
    cells = [norm(merged_cell_value(ws, row, c)) for c in range(1, max_col + 1)]
    REQUIRED = {
        "staff": ["שם המרצה", "מרצה"],
        "group": ["קבוצת מעבדה", "קבוצה"],
        "time": ["שעה"],
        "day": ["יום"],
        "date": ["תאריך"],
        "sessionNo": ["מס' מע'", "מספר מעבדה", "מס׳ מע", "מס' מעבדה"],
    }
    mapping = {}
    hits = 0
    for key, variants in REQUIRED.items():
        for i, txt in enumerate(cells, start=1):
            if txt and any(v in txt for v in variants):
                mapping[key] = i
                hits += 1
                break
    return mapping, hits


def extract_course_from_line(text):
    t = norm(text)
    if not t:
        return None, None
    m = re.search(r"(.+?)\s*[-–]\s*(\d{4,6})$", t)
    if m:
        return m.group(2), norm(m.group(1))
    m = re.search(r"^(\d{4,6})\s*[-–]\s*(.+)$", t)
    if m:
        return m.group(1), norm(m.group(2))
    return None, None


def find_course_title_near(ws, header_row, max_col):
    for r in range(header_row - 1, max(1, header_row - 8), -1):
        line = " ".join(
            norm(merged_cell_value(ws, r, c))
            for c in range(1, max_col + 1)
            if merged_cell_value(ws, r, c)
        )
        code, name = extract_course_from_line(line)
        if code and name:
            return code, name
    return None, None


def parse_labs_from_excel(path: Path, semester: int) -> list[dict]:
    wb = load_workbook(path, data_only=True)
    rows = []
    for ws in wb.worksheets:
        r = 1
        while r <= ws.max_row:
            header, hits = build_lab_header(ws, r, ws.max_column)
            if hits < 5:
                r += 1
                continue
            course_code, course_name = find_course_title_near(ws, r, ws.max_column)
            if not course_code:
                r += 1
                continue
            rr = r + 1
            while rr <= ws.max_row:
                raw_date = merged_cell_value(ws, rr, header.get("date", 0) or 1)
                raw_day = merged_cell_value(ws, rr, header.get("day", 0) or 1)
                raw_staff = merged_cell_value(ws, rr, header.get("staff", 0) or 1) if header.get("staff") else None
                raw_time = merged_cell_value(ws, rr, header.get("time", 0) or 1) if header.get("time") else None
                raw_group = merged_cell_value(ws, rr, header.get("group", 0) or 1) if header.get("group") else None

                session_val = ""
                if header.get("sessionNo"):
                    session_val = norm(merged_cell_value(ws, rr, header["sessionNo"]))

                date_val = norm(raw_date)
                day_val = norm(raw_day)
                if not session_val and not date_val and not day_val:
                    break

                # Also read WITHOUT merged-cell helper (parser behavior)
                plain_date = norm(ws.cell(rr, header.get("date", 0)).value) if header.get("date") else ""
                plain_staff = norm(ws.cell(rr, header["staff"]).value) if header.get("staff") else ""

                # Fill-down from previous row in sheet
                prev_date = norm(merged_cell_value(ws, rr - 1, header["date"])) if header.get("date") and rr > 1 else ""
                prev_staff = norm(merged_cell_value(ws, rr - 1, header["staff"])) if header.get("staff") and rr > 1 else ""

                rows.append({
                    "semester": semester,
                    "sourceFile": path.name,
                    "courseCode": str(course_code),
                    "courseName": course_name,
                    "session": session_val,
                    "row": rr,
                    "group": norm(raw_group),
                    "dateMerged": date_val,
                    "datePlain": plain_date,
                    "staffMerged": norm(raw_staff),
                    "staffPlain": plain_staff,
                    "time": norm(raw_time),
                    "day": day_val,
                    "prevRowDate": prev_date,
                    "prevRowStaff": prev_staff,
                    "isMergedDateCell": plain_date == "" and date_val != "",
                    "isMergedStaffCell": plain_staff == "" and norm(raw_staff) != "",
                    "fillDownWouldFixDate": plain_date == "" and prev_date != "",
                    "fillDownWouldFixStaff": plain_staff == "" and prev_staff != "",
                })
                rr += 1
            r = rr
    return rows


def fetch_firebase_labs(db) -> list[dict]:
    rows = []
    for year in db.collection("lab_schedule").stream():
        year_id = year.id
        for sem in year.reference.collection("semesters").stream():
            sem_data = sem.to_dict() or {}
            sem_num = sem_data.get("semester") or sem.id
            courses = sem_data.get("courses") or {}
            for code, course in courses.items():
                for i, lab in enumerate(course.get("labs") or []):
                    staff = lab.get("staff") or []
                    rows.append({
                        "yearId": year_id,
                        "semester": sem_num,
                        "courseCode": code,
                        "courseName": course.get("courseName"),
                        "labIndex": i,
                        "session": lab.get("session"),
                        "date": lab.get("date"),
                        "day": lab.get("day"),
                        "time": lab.get("time"),
                        "group": lab.get("group"),
                        "staff": staff,
                        "missing": [
                            f for f in ("date", "day", "time", "group", "staff")
                            if (
                                (f == "staff" and (not staff or not len(staff)))
                                or (f != "staff" and not lab.get(f))
                            )
                        ],
                    })
    return rows


def discover_lab_files() -> list[tuple[Path, int]]:
    labs_dir = SCRIPT_DIR / "labs"
    found = []
    for path in sorted(labs_dir.glob("*.xlsx")):
        m = LAB_SEMESTER_RE.search(path.name)
        if m:
            found.append((path, int(m.group(1))))
    return found


def analyze_missing_labs(db) -> dict:
    fb_labs = fetch_firebase_labs(db)
    fb_missing = [r for r in fb_labs if r["missing"]]

    excel_rows: list[dict] = []
    for path, sem in discover_lab_files():
        excel_rows.extend(parse_labs_from_excel(path, sem))

    # Index excel by semester + course + session + group + row approx
    report = []
    for fb in fb_missing:
        sem = int(fb["semester"]) if str(fb["semester"]).isdigit() else fb["semester"]
        candidates = [
            e for e in excel_rows
            if e["semester"] == sem
            and e["courseCode"] == fb["courseCode"]
            and (e["session"] == fb["session"] or e["group"] == fb["group"])
        ]
        best = candidates[0] if candidates else None

        diagnosis = {}
        for field in fb["missing"]:
            if field == "staff":
                if best:
                    if best["staffPlain"]:
                        diagnosis["staff"] = "parser_gap_plain_cell_has_value"
                    elif best["staffMerged"]:
                        diagnosis["staff"] = "parser_gap_merged_cell_has_value"
                    elif best["fillDownWouldFixStaff"]:
                        diagnosis["staff"] = "excel_empty_fill_down_would_help"
                    else:
                        diagnosis["staff"] = "genuinely_missing_in_excel"
                else:
                    diagnosis["staff"] = "no_matching_excel_row"
            elif field == "date":
                if best:
                    if best["datePlain"]:
                        diagnosis["date"] = "parser_gap_plain_cell_has_value"
                    elif best["dateMerged"]:
                        diagnosis["date"] = "parser_gap_merged_cell_has_value"
                    elif best["fillDownWouldFixDate"]:
                        diagnosis["date"] = "excel_empty_fill_down_would_help"
                    else:
                        diagnosis["date"] = "genuinely_missing_in_excel"
                else:
                    diagnosis["date"] = "no_matching_excel_row"
            else:
                if best and best.get(field):
                    diagnosis[field] = "parser_gap"
                elif best:
                    diagnosis[field] = "genuinely_missing_in_excel"
                else:
                    diagnosis[field] = "no_matching_excel_row"

        report.append({
            **{k: fb[k] for k in ("yearId", "semester", "courseCode", "courseName", "session", "group", "date", "staff", "missing")},
            "excelMatch": best,
            "fieldDiagnosis": diagnosis,
        })

    missing_date = sum(1 for r in fb_missing if "date" in r["missing"])
    missing_staff = sum(1 for r in fb_missing if "staff" in r["missing"])
    parser_fixable = sum(
        1 for r in report
        if any(v.endswith("_has_value") or v == "parser_gap" or "fill_down" in v for v in r["fieldDiagnosis"].values())
    )

    return {
        "firebaseLabRows": len(fb_labs),
        "firebaseRowsWithMissingFields": len(fb_missing),
        "missingDateCount": missing_date,
        "missingStaffCount": missing_staff,
        "parserFixableEstimate": parser_fixable,
        "rows": report,
    }


def analyze_cred_column_issue() -> dict:
    """Check if find_col('נ') falsely matches unrelated headers."""
    samples = []
    for yb_id, path in YEARBOOK_SOURCES.items():
        if not path.is_file():
            continue
        doc = Document(str(path))
        table_map = {t._element: t for t in doc.tables}
        current_sem = None
        for block in doc.element.body:
            if block.tag.endswith("p"):
                text = normalize("".join(t.text for t in block.xpath(".//w:t")))
                m = re.search(r"סמסטר\s*([1-8])", text)
                if m:
                    current_sem = int(m.group(1))
            elif block.tag.endswith("tbl") and current_sem:
                table = table_map.get(block)
                if not table or not table.rows:
                    continue
                headers = [normalize(c.text) for c in table.rows[0].cells]
                if "שם הקורס" not in " ".join(headers):
                    continue
                cred_i = find_col(headers, 'נ"ז', "נקודות", "נ")
                samples.append({
                    "yearbookId": yb_id,
                    "semester": current_sem,
                    "headers": headers,
                    "credColIndex": cred_i,
                    "credHeader": headers[cred_i] if cred_i is not None else None,
                })
                break  # one table per file enough
    return {"headerSamples": samples}


def json_default(obj):
    if isinstance(obj, (datetime, date, time)):
        return obj.isoformat()
    return str(obj)


def main():
    db = init_firebase()

    report = {
        "mode": "READ_ONLY_QUALITY_CHECK",
        "testYearbook": inspect_test_yearbook(db),
        "missingCredits": analyze_missing_credits(db),
        "credColumnAnalysis": analyze_cred_column_issue(),
        "missingLabs": analyze_missing_labs(db),
        "recommendations": [],
    }

    # Build recommendations
    t = report["testYearbook"]
    if t.get("deleteRecommendation") in ("safe_to_delete_after_approval", "likely_stale_shell_delete_after_approval"):
        report["recommendations"].append(
            "Delete yearbooks/test after manual approval — empty shell with no requiredCourses/courses."
        )

    mc = report["missingCredits"]
    parser_gaps = [r for r in mc["rows"] if "parser" in r["diagnosis"] or r["diagnosis"] == "parser_failed_non_numeric_source"]
    genuine = [r for r in mc["rows"] if r["diagnosis"] == "genuinely_missing_in_source_docx"]
    if parser_gaps:
        report["recommendations"].append(
            f"Credits: {len(parser_gaps)} row(s) may be parser issues — review cred column detection (needle 'נ' may match wrong column)."
        )
    if genuine:
        report["recommendations"].append(
            f"Credits: {len(set(r['courseCode'] for r in genuine))} unique course code(s) genuinely empty in source DOCX — not a parser bug."
        )

    ml = report["missingLabs"]
    if ml["parserFixableEstimate"]:
        report["recommendations"].append(
            "Labs: add merged-cell reading and/or fill-down for date/staff in labs_parser.py, then re-import lab semesters only."
        )
    if ml["missingDateCount"] or ml["missingStaffCount"]:
        report["recommendations"].append(
            f"Labs: {ml['missingDateCount']} rows missing date, {ml['missingStaffCount']} missing staff in Firebase."
        )

    report["recommendations"].append(
        "After parser fixes: run import_to_firebase.py --write for labs only (or full re-import) — requires explicit approval."
    )

    OUTPUT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2, default=json_default), encoding="utf-8")

    # Console summary
    print("=" * 60)
    print("POST-IMPORT QUALITY CHECK (read-only)")
    print("=" * 60)
    print(f"\nReport written: {OUTPUT_PATH}")

    print("\n--- yearbooks/test ---")
    print(json.dumps(t, ensure_ascii=False, indent=2))

    print(f"\n--- Missing credits ---")
    print(f"  Firebase rows missing credits (excl. test): {mc['firebaseMissingCreditsCount']}")
    print(f"  Unique course codes: {mc['uniqueCourseCodesMissingCredits']}")
    for u in mc["uniqueSummary"][:20]:
        print(f"    {u['courseCode']} {u['courseName']}: {u['diagnoses']} (in {len(u['yearbooks'])} yearbooks)")

    print(f"\n--- Missing labs ---")
    print(f"  Rows with missing fields: {ml['firebaseRowsWithMissingFields']} / {ml['firebaseLabRows']}")
    print(f"  Missing date: {ml['missingDateCount']}, missing staff: {ml['missingStaffCount']}")
    print(f"  Parser-fixable estimate: {ml['parserFixableEstimate']}")

    print("\n--- Recommendations ---")
    for r in report["recommendations"]:
        print(f"  * {r}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
