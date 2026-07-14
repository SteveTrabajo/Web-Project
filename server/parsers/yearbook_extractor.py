import sys
import os
import re
import json

"""
yearbook_extractor.py

Raw content extractor for yearbook files (DOCX or PDF). This does NOT write to
Firestore and does NOT decide course relations - it only pulls the semester
course tables out into a normalized JSON structure. The Node layer sends that
structure to the LLM for structured extraction (see services/yearbookImport.js).

Only the per-semester course tables are extracted. Prose, intro pages, footnotes,
and the elective / specialization ("לימודי התמחות", "קורסי בחירה") sections are
skipped - the importer targets the required curriculum (semesters 1-8) only.

Output: a single JSON object printed to stdout:
{
  "format": "docx" | "pdf",
  "semesters": [
    { "semesterNumber": 1, "headers": [...], "rows": [[cell, cell, ...], ...] }
  ],
  "looseTables": [ { "headers": [...], "rows": [...] } ],  # course tables with no detectable semester
  "warnings": [ "..." ]
}

Usage: python yearbook_extractor.py <file_path>
"""

SEM_RE = re.compile(r"סמסטר\s*([1-8])")
CODE_RE = re.compile(r"^\d{5,6}$")

# A page count over this is almost certainly a full yearbook, not the course
# section - reject it and ask the admin to upload only the course-table pages.
MAX_PDF_PAGES = 20
# DOCX has no pages; cap on table count instead as a sanity guard.
MAX_DOCX_TABLES = 80

# Start of the electives / specialization section. Once seen, the remaining
# tables are track courses and choice lists (no semester) - stop importing.
# Matched against whitespace-stripped text: PDF extraction routinely injects
# spurious spaces inside Hebrew words (e.g. "התמחו ת" for "התמחות").
ELECTIVES_RE = re.compile(r"לימודיהתמחות|קורסיבחירה|לימודיבחירה|לימודיהשלמה")


def is_electives_marker(text):
    return bool(ELECTIVES_RE.search(re.sub(r"\s+", "", text or "")))


def is_semester_heading(text):
    """A short standalone 'סמסטר N' heading (not prose that merely mentions it)."""
    m = SEM_RE.search(text)
    return int(m.group(1)) if (m and len(text.strip()) <= 15) else None


def normalize(s):
    if not s:
        return ""
    return re.sub(r"\s+", " ", s.replace(" ", " ")).strip()


def looks_like_course_table(rows):
    """A real course table has >=5 columns and at least one row whose cells
    include a bare 5-6 digit course code. The column count already excludes the
    single-column prose/note tables; the code check confirms it holds courses."""
    if not rows or len(rows[0]) < 5:
        return False
    return any(
        CODE_RE.match((c or "").strip()) for row in rows for c in row
    )


# ==============================
# DOCX extraction
# ==============================
def render_cell_docx(cell):
    """Render a DOCX cell to text, wrapping underlined runs in <u>...</u> so the
    corequisite signal survives into the plain-text sent to the LLM."""
    parts = []
    for p in cell.paragraphs:
        line = []
        for run in p.runs:
            text = run.text
            if not text:
                continue
            if run.font.underline:
                line.append(f"<u>{text}</u>")
            else:
                line.append(text)
        para = normalize("".join(line))
        if para:
            parts.append(para)
    return " | ".join(parts)


def extract_docx(file_path):
    from docx import Document

    doc = Document(file_path)
    if len(doc.tables) > MAX_DOCX_TABLES:
        raise ValueError(
            f"Document has {len(doc.tables)} tables - too large. Upload only the course-table pages."
        )

    table_map = {t._element: t for t in doc.tables}

    semesters = []
    loose_tables = []
    warnings = []
    current_sem = None
    in_electives = False
    skipped_electives = 0
    sem_bucket = {}

    def bucket_for(sem):
        if sem not in sem_bucket:
            entry = {"semesterNumber": sem, "headers": [], "rows": []}
            sem_bucket[sem] = entry
            semesters.append(entry)
        return sem_bucket[sem]

    for block in doc.element.body:
        if block.tag.endswith("p"):
            text = normalize("".join(t.text for t in block.xpath(".//w:t")))
            if is_electives_marker(text):
                in_electives = True
            sem = is_semester_heading(text)
            if sem is not None:
                current_sem = sem
                in_electives = False
        elif block.tag.endswith("tbl"):
            table = table_map.get(block)
            if not table or not table.rows:
                continue
            headers = [normalize(c.text) for c in table.rows[0].cells]
            rows = [[render_cell_docx(c) for c in row.cells] for row in table.rows[1:]]
            if not looks_like_course_table([headers] + rows):
                continue
            if in_electives:
                skipped_electives += 1
                continue
            if current_sem is not None:
                entry = bucket_for(current_sem)
                if not entry["headers"]:
                    entry["headers"] = headers
                entry["rows"].extend(rows)
            else:
                loose_tables.append({"headers": headers, "rows": rows})

    return _finish(semesters, loose_tables, skipped_electives, warnings, "docx")


# ==============================
# PDF extraction
# ==============================
def extract_pdf(file_path):
    import pdfplumber
    from bidi.algorithm import get_display

    # pdfplumber/pdfminer return RTL text in visual (reversed) order. get_display
    # converts it back to logical order per line, handling embedded numbers.
    def fix_rtl(s):
        if not s:
            return s
        return "\n".join(get_display(ln) for ln in s.split("\n"))

    semesters = []
    loose_tables = []
    warnings = []
    skipped_electives = 0
    sem_bucket = {}

    def bucket_for(sem):
        if sem not in sem_bucket:
            entry = {"semesterNumber": sem, "headers": [], "rows": []}
            sem_bucket[sem] = entry
            semesters.append(entry)
        return sem_bucket[sem]

    def clean_table(data):
        out = []
        for row in data or []:
            cells = [normalize(fix_rtl(c or "")) for c in row]
            if any(c for c in cells):
                out.append(cells)
        return out

    with pdfplumber.open(file_path) as pdf:
        if len(pdf.pages) > MAX_PDF_PAGES:
            raise ValueError(
                f"PDF has {len(pdf.pages)} pages - the limit is {MAX_PDF_PAGES}. "
                "Upload only the course-table pages."
            )

        current_sem = None
        in_electives = False
        for page in pdf.pages:
            # Semester headings on this page, with their vertical position. A page
            # can hold several semesters, so each table is matched to the nearest
            # heading above it (not just the last heading on the page).
            headings = []
            for ln in page.extract_text_lines():
                t = fix_rtl(ln.get("text", "")).strip()
                if is_electives_marker(t):
                    in_electives = True
                sem = is_semester_heading(t)
                if sem is not None:
                    headings.append((ln["top"], sem))
                    in_electives = False
            headings.sort()

            for tb in sorted(page.find_tables(), key=lambda t: t.bbox[1]):
                rows = clean_table(tb.extract())
                if not looks_like_course_table(rows):
                    continue
                if in_electives:
                    skipped_electives += 1
                    continue

                # Nearest semester heading above this table; else carry over.
                top = tb.bbox[1]
                sem = current_sem
                for htop, hsem in headings:
                    if htop <= top:
                        sem = hsem

                headers = rows[0]
                body = rows[1:]
                if not body:
                    continue
                if sem is not None:
                    entry = bucket_for(sem)
                    if not entry["headers"]:
                        entry["headers"] = headers
                    entry["rows"].extend(body)
                else:
                    loose_tables.append({"headers": headers, "rows": body})

            # Carry the last (lowest) heading on this page into the next page so a
            # table continuing across a page break keeps its semester.
            if headings:
                current_sem = headings[-1][1]

    return _finish(semesters, loose_tables, skipped_electives, warnings, "pdf")


def _finish(semesters, loose_tables, skipped_electives, warnings, fmt):
    if not semesters and not loose_tables:
        warnings.append("No course tables detected (the file may be scanned/image-based or contain no course tables).")
    if loose_tables:
        warnings.append(f"{len(loose_tables)} course table(s) had no detectable semester heading.")
    if skipped_electives:
        warnings.append(f"{skipped_electives} elective/specialization table(s) were skipped - only semester courses are imported.")
    return {"format": fmt, "semesters": semesters, "looseTables": loose_tables, "warnings": warnings}


# ==============================
# Entry point
# ==============================
def main():
    # Windows consoles default to a non-UTF-8 codepage (cp1251/cp1252), which
    # crashes on Hebrew output. Force UTF-8 where supported.
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: yearbook_extractor.py <file_path>"}))
        sys.exit(1)

    file_path = sys.argv[1]
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".docx":
            result = extract_docx(file_path)
        elif ext == ".pdf":
            result = extract_pdf(file_path)
        else:
            print(json.dumps({"error": f"Unsupported file type: {ext}. Use .docx or .pdf"}))
            sys.exit(1)
    except Exception as e:
        print(json.dumps({"error": f"Extraction failed: {e}"}))
        sys.exit(1)

    # ensure_ascii=True keeps output pure ASCII (Hebrew as \uXXXX), which Node's
    # JSON.parse handles fine and which never hits a console-encoding error.
    print(json.dumps(result))


if __name__ == "__main__":
    main()
