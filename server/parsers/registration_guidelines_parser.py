import sys
import re
import json
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

"""
registration_guidelines_parser.py

קורא קובץ DOCX של "הנחיות כלליות לרישום" ומחזיר טיוטה במבנה של
registrationGuidelines/semester_N (ללא כתיבה ל-Firestore - המנהל מאשר בממשק).

- טקסט חופשי -> keyRules (כלל אחד לפסקה, עם קוד פנימי לפי נושא)
- טבלת היועצים -> contacts לפי קטגוריה (יועצים, פטורים, מעבדות, מלווה, מזכירות)
- מספרי נ"ז (165 לתואר, מינימום/מקסימום סמסטריאלי) -> audience
- קישורים -> links

שימוש:
python registration_guidelines_parser.py <file_path>
"""

# ==============================
# Helpers
# ==============================
HEB_LETTERS = "אבגדהוזחטיכלמנסעפצקרשת"

# Track names must match TRACKS in Bot.jsx / AdvisorsTab.jsx.
TRACK_MOLECULAR = "מולקולרית-תרופתית"
TRACK_FOOD_ENV = "מזון והסביבה"

# Advisors on a specialization track only apply from semester 5 onward
# (advisor.js ignores `track` below semester 5).
TRACK_SEMESTERS = [5, 6, 7, 8]


def norm(x):
    """Collapse whitespace and strip separator noise left by the DOCX cells."""
    if x is None:
        return ""
    s = re.sub(r"\s+", " ", str(x).replace("\n", " "))
    return s.strip()


def clean_cell(x):
    """Cell text without the trailing '/' and dash fragments the source is full of."""
    s = norm(x)
    s = re.sub(r"[\s/\-–—]+$", "", s)
    s = re.sub(r"^[\s/\-–—]+", "", s)
    # Unbalanced parenthesis left over from the source (e.g. "0-12.99 (")
    if s.count("(") != s.count(")"):
        s = s.replace("(", " ").replace(")", " ")
    return norm(s)


def iter_blocks(doc):
    """Paragraphs and tables in document order.

    python-docx exposes doc.paragraphs and doc.tables as separate lists, which
    loses the interleaving - and here the advisor table sits between two prose
    sections, so order carries meaning.
    """
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def row_texts(row, width):
    """Row cells as a positional list of length `width`.

    A horizontally merged cell is returned by python-docx once per grid column
    it spans. Keeping the text on the first position and blanking the repeats
    preserves column alignment instead of shifting later columns left.
    """
    out = []
    seen = set()
    for cell in row.cells:
        key = id(cell._tc)
        out.append("" if key in seen else norm(cell.text))
        seen.add(key)
    out = out[:width]
    while len(out) < width:
        out.append("")
    return out


# ==============================
# Field extraction
# ==============================
def parse_email(text):
    m = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text or "")
    return m.group(0).strip(" .,;") if m else ""


def parse_semesters(text):
    """'סמסטר 1+2' -> [1,2] ; 'סמסטר 5-8' -> [5,6,7,8] ; 'סמסטר 3' -> [3]."""
    t = text or ""
    m = re.search(r"סמסטר\s*(\d)\s*\+\s*(\d)", t)
    if m:
        return sorted({int(m.group(1)), int(m.group(2))})
    m = re.search(r"סמסטר\s*(\d)\s*[-–]\s*(\d)", t)
    if m:
        a, b = int(m.group(1)), int(m.group(2))
        return list(range(min(a, b), max(a, b) + 1))
    m = re.search(r"סמסטר\s*(\d)", t)
    if m:
        return [int(m.group(1))]
    return []


def parse_credits_range(text):
    """'טווח נ\"ז 0-12.99' -> {min: 0.0, max: 12.99}."""
    m = re.search(r"(\d+(?:\.\d+)?)\s*[-–]\s*(\d+(?:\.\d+)?)", text or "")
    if not m:
        return None
    a, b = float(m.group(1)), float(m.group(2))
    return {"min": min(a, b), "max": max(a, b)}


def parse_last_name_range(text):
    """Surname bucket out of the notes column.

    The source phrases it inconsistently: "שם משפחה -( (א' עד כ')", "(ל' עד ת')",
    "א-כ". All collapse to a from/to pair of single Hebrew letters.
    """
    t = text or ""
    m = re.search(r"([{L}])\s*['׳’]?\s*(?:עד|[-–])\s*([{L}])".format(L=HEB_LETTERS), t)
    if not m:
        return None, None
    return m.group(1), m.group(2)


def parse_effective_from(text):
    """'מתאריך 30.3.26' -> '30.3.26'. Real validity data, not a footnote."""
    m = re.search(r"מתאריך\s*(\d{1,2}[./]\d{1,2}[./]\d{2,4})", text or "")
    return m.group(1) if m else ""


def classify_row(topic):
    """Map the advisory-topic column onto a registrationGuidelines contact bucket."""
    t = topic or ""
    if "פטור" in t:
        return "exemptions"
    if "מלוו" in t or "מלווה" in t:
        return "mentors"
    if "מעבד" in t:
        return "labs"
    if "רמ" in t and "ח" in t:
        return "registrationSupport"
    return "academicAdvisors"


def parse_track(topic):
    t = topic or ""
    if "מולקולרית" in t or "תרופתית" in t:
        return TRACK_MOLECULAR
    if "מזון" in t or "סביבה" in t:
        return TRACK_FOOD_ENV
    return ""


# ==============================
# Prose -> keyRules
# ==============================
RULE_CODES = [
    ("INTERNSHIP", ("סטאז", "התמחות מעשית")),
    ("CREDITS", ("נ\"ז", "נ״ז", "נקודות זכות", "נקודות הזכות")),
    ("EXAMS", ("בחינ", "מבחנ", "לוח המבחנים")),
    ("LABS", ("מעבד",)),
    ("PREREQ", ("תנאי קדם", "תנאי הקדם", "קדם")),
    ("SYLLABUS", ("סילבוס",)),
    ("SCHEDULE", ("מערכת השעות", "מערכת שעות")),
    ("REGISTRATION", ("רישום",)),
]


def rule_code(text, idx):
    for code, needles in RULE_CODES:
        if any(n in text for n in needles):
            return "{}_{}".format(code, idx)
    return "GEN_{}".format(idx)


def is_section_header(text):
    """Short line ending in ':' or '-' that labels the block below it."""
    return len(text) <= 30 and bool(re.search(r"[-:–]\s*$", text))


# ==============================
# Main parse
# ==============================
def parse(path):
    doc = Document(str(path))
    warnings = []

    titles = []
    prose = []
    advisor_table = None

    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            text = norm(block.text)
            if not text:
                continue
            style = (block.style.name or "").lower()
            if "title" in style or "heading" in style:
                titles.append(text)
            else:
                prose.append(text)
        else:
            header = " ".join(row_texts(block.rows[0], len(block.columns))) if block.rows else ""
            if advisor_table is None and ("יועץ" in header or "שם היועץ" in header):
                advisor_table = block
            elif advisor_table is None and block.rows:
                # Only one table is expected; keep the first as a fallback.
                advisor_table = block

    if advisor_table is None:
        raise Exception("no advisor table found in document")

    # ---- meta ----
    joined_titles = " ".join(titles)
    year = ""
    m = re.search(r"(תשפ[\"״']?[{L}])".format(L=HEB_LETTERS), joined_titles)
    if m:
        year = m.group(1)
    term = ""
    if re.search(r"סמסטר\s*א", joined_titles):
        term = "A"
    elif re.search(r"סמסטר\s*ב", joined_titles):
        term = "B"

    title = next((t for t in titles if "הנחיות" in t), joined_titles.strip(" -"))

    # ---- credits ----
    all_prose = " ".join(prose)
    # A degree total is always three digits; a semester cap is two. Matching on
    # magnitude avoids depending on which sentence happens to come first.
    degree_total = None
    for m in re.finditer(r"(\d{2,3})\s*נ[\"״']?ז", all_prose):
        n = int(m.group(1))
        if 100 <= n <= 400:
            degree_total = n
            break

    sem_max = sem_min = None
    m = re.search(r"מקסימלי\S*\s+בסמסטר\s+ה[ןם]\s*(\d+)", all_prose)
    if m:
        sem_max = int(m.group(1))
    m = re.search(r"מינימלי\S*\s+ה[ןם]\s*(\d+)", all_prose)
    if m:
        sem_min = int(m.group(1))
    if sem_max is None or sem_min is None:
        warnings.append("לא זוהו מגבלות נ\"ז סמסטריאליות (מינימום/מקסימום) בטקסט")

    credits_rule_text = ""
    for p in prose:
        if "מקסימלי" in p and "נ" in p:
            credits_rule_text = p
            break

    # ---- links ----
    links = []
    seen_urls = set()
    for p in prose:
        for url in re.findall(r"https?://[^\sא-ת,)]+", p):
            url = url.rstrip(".,;")
            if url not in seen_urls:
                seen_urls.add(url)
                links.append({"label": url, "url": url})
    for rel in doc.part.rels.values():
        if rel.reltype.endswith("/hyperlink"):
            url = str(rel.target_ref).rstrip(".,;")
            if url.startswith("http") and url not in seen_urls:
                seen_urls.add(url)
                links.append({"label": url, "url": url})

    # ---- key rules ----
    key_rules = []
    idx = 0
    for p in prose:
        if is_section_header(p):
            continue
        idx += 1
        key_rules.append({"code": rule_code(p, idx), "text": p})

    # ---- advisor table ----
    width = len(advisor_table.columns)
    header = row_texts(advisor_table.rows[0], width)

    def col_index(*needles):
        for i, h in enumerate(header):
            if any(n in h for n in needles):
                return i
        return None

    i_topic = col_index("נושא", "טווח") or 0
    i_name = col_index("שם היועץ", "שם")
    i_email = col_index("מייל", "דוא")
    i_notes = col_index("הערות")
    if i_name is None:
        i_name = 1
    if i_email is None:
        i_email = 2
    if i_notes is None:
        i_notes = 3

    contacts = {
        "registrationSupport": [],
        "mentors": [],
        "academicAdvisors": [],
        "exemptions": [],
        "labs": [],
    }
    rows_out = []

    for r in advisor_table.rows[1:]:
        cells = row_texts(r, width)
        topic = clean_cell(cells[i_topic] if i_topic < width else "")
        name = clean_cell(cells[i_name] if i_name < width else "")
        notes = norm(cells[i_notes] if i_notes < width else "")
        email = parse_email(cells[i_email] if i_email < width else "") or parse_email(notes)

        if not topic and not name:
            continue

        category = classify_row(topic)
        track = parse_track(topic)
        semesters = parse_semesters(topic)
        if not semesters and track:
            semesters = list(TRACK_SEMESTERS)
        credits_range = parse_credits_range(topic)
        ln_from, ln_to = parse_last_name_range(notes)
        effective_from = parse_effective_from(notes)

        if not name:
            warnings.append('שורה בטבלה ללא שם ממלא תפקיד: "{}"'.format(topic))
        if not email:
            warnings.append('אין מייל עבור "{}"{}'.format(name or topic, " ({})".format(topic) if name else ""))
        if category == "academicAdvisors" and not track and not ln_from:
            warnings.append('לא זוהה שיוך אלפביתי עבור "{}"'.format(name or topic))

        row = {
            "category": category,
            "topic": topic,
            "name": name,
            "email": email,
            "notes": notes,
            "semesters": semesters,
            "creditsRange": credits_range,
            "track": track,
            "lastNameFrom": ln_from or "",
            "lastNameTo": ln_to or "",
            "effectiveFrom": effective_from,
        }
        rows_out.append(row)

        if category == "academicAdvisors":
            contacts["academicAdvisors"].append({
                "name": name,
                "email": email,
                "assignment": {
                    "lastNameFrom": ln_from or "",
                    "lastNameTo": ln_to or "",
                    "track": track,
                },
                "semesters": semesters,
                "creditsRange": credits_range,
                "effectiveFrom": effective_from,
            })
        elif category == "labs":
            contacts["labs"].append({
                "name": name,
                "role": topic,
                "email": email,
                "howToContact": notes,
                "semesters": semesters,
            })
        elif category == "mentors":
            contacts["mentors"].append({
                "name": name,
                "role": topic,
                "email": email,
                "semesters": semesters,
            })
        else:
            contacts[category].append({
                "name": name,
                "role": topic,
                "email": email,
                "phone": "",
                "semesters": semesters,
            })

    return {
        "ok": True,
        "meta": {
            "title": title,
            "year": year,
            "term": term,
            "termText": joined_titles.strip(" -"),
        },
        "audience": {
            "cohortText": " ".join(x for x in [year, "סמסטר א'" if term == "A" else ("סמסטר ב'" if term == "B" else "")] if x).strip(),
            "creditsRuleText": credits_rule_text or None,
            "creditsRange": ({"min": sem_min, "max": sem_max} if sem_min is not None or sem_max is not None else None),
            "degreeCredits": degree_total,
        },
        "keyRules": key_rules,
        "links": links,
        "contacts": contacts,
        "advisorRows": rows_out,
        "warnings": warnings,
    }


# ==============================
# ENTRY POINT
# ==============================
if __name__ == "__main__":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    if len(args) < 1:
        raise Exception("Usage: registration_guidelines_parser.py <file_path>")

    result = parse(Path(args[0]))
    # Last stdout line is the JSON payload consumed by uploadAdmin.js
    print(json.dumps(result, ensure_ascii=False))
